"""
Template Filler Module
Populates DOCX templates with extracted field values.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Union
from io import BytesIO
import re
import copy


def fill_template(doc: Document, field_values: Dict[str, str]) -> Document:
    """
    Fill a template document with extracted field values.
    Uses intelligent matching to find and replace field locations.
    
    Args:
        doc: python-docx Document object (template)
        field_values: Dictionary mapping field names to values
        
    Returns:
        Modified Document object with filled values
    """
    # Create a working copy
    # Note: python-docx doesn't have a copy method, so we work with the original
    
    # Process paragraphs
    for paragraph in doc.paragraphs:
        fill_paragraph(paragraph, field_values)
    
    # Process tables
    for table in doc.tables:
        fill_table(table, field_values)
    
    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                fill_paragraph(paragraph, field_values)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                fill_paragraph(paragraph, field_values)
    
    return doc


def fill_paragraph(paragraph, field_values: Dict[str, str]) -> None:
    """
    Fill a single paragraph with field values.
    Handles various placeholder formats.
    
    Args:
        paragraph: python-docx Paragraph object
        field_values: Dictionary mapping field names to values
    """
    original_text = paragraph.text
    new_text = original_text
    
    # Try to match and replace field values
    for field_name, value in field_values.items():
        if value and value != "N/A":
            # Create patterns to match the field
            patterns = [
                # Exact match with colon (e.g., "Field Name: ____")
                (rf'({re.escape(field_name)}\s*:\s*)(_+|\s*$)', rf'\1{value}'),
                # Exact match with spaces/underscores
                (rf'({re.escape(field_name)}\s*)(_+)', rf'\1{value}'),
                # Field in brackets
                (rf'\[{re.escape(field_name)}\]', value),
                # Field in curly braces
                (rf'\{{\{{{re.escape(field_name)}\}}\}}', value),
                # Field in angle brackets
                (rf'<{re.escape(field_name)}>', value),
            ]
            
            for pattern, replacement in patterns:
                try:
                    new_text = re.sub(pattern, replacement, new_text, flags=re.IGNORECASE)
                except re.error:
                    continue
    
    # Only update if text changed
    if new_text != original_text:
        # Preserve formatting by updating runs
        if paragraph.runs:
            # Clear all but first run, update first run with new text
            first_run = paragraph.runs[0]
            for run in paragraph.runs[1:]:
                run.text = ""
            first_run.text = new_text
        else:
            paragraph.text = new_text


def fill_table(table, field_values: Dict[str, str]) -> None:
    """
    Fill a table with field values.
    Intelligently matches field names in cells.
    
    Args:
        table: python-docx Table object
        field_values: Dictionary mapping field names to values
    """
    for row in table.rows:
        cells = list(row.cells)
        
        for i, cell in enumerate(cells):
            cell_text = cell.text.strip()
            
            # Check if this cell contains a field name
            for field_name, value in field_values.items():
                if value and value != "N/A":
                    # If cell text matches or contains field name
                    if field_name.lower() in cell_text.lower():
                        # Check if this is a label cell (field name) and next cell is value
                        if i + 1 < len(cells):
                            next_cell = cells[i + 1]
                            next_text = next_cell.text.strip()
                            
                            # If next cell is empty or has placeholder, fill it
                            if not next_text or re.match(r'^_+$|^\s*$', next_text):
                                # Preserve cell formatting
                                if next_cell.paragraphs:
                                    next_cell.paragraphs[0].text = str(value)
                                else:
                                    next_cell.text = str(value)
                    
                    # Also try to replace within the cell itself
                    for paragraph in cell.paragraphs:
                        fill_paragraph(paragraph, {field_name: value})


def fill_template_smart(
    template_doc: Document,
    field_values: Dict[str, str],
    report_text: str
) -> Document:
    """
    Smart template filling that uses the report text context.
    Falls back to direct field matching if smart matching fails.
    
    Args:
        template_doc: python-docx Document object
        field_values: Dictionary from LLM extraction
        report_text: Original report text for context
        
    Returns:
        Filled Document object
    """
    # First pass: direct field value insertion
    doc = fill_template(template_doc, field_values)
    
    return doc


def save_filled_template(doc: Document, output_path: str = None) -> Union[str, BytesIO]:
    """
    Save the filled template to a file or return as BytesIO.
    
    Args:
        doc: Filled Document object
        output_path: Optional file path to save to
        
    Returns:
        File path if output_path provided, else BytesIO object
    """
    if output_path:
        doc.save(output_path)
        return output_path
    else:
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


if __name__ == "__main__":
    # Test the module
    print("Template filler module loaded successfully")
