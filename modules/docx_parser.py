"""
DOCX Template Parser Module
Parses and analyzes insurance templates to identify structure and content.
"""

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from typing import Union, List, Dict
from io import BytesIO
import re


def parse_template(docx_source: Union[str, bytes, BytesIO]) -> Document:
    """
    Load and parse a DOCX template file.
    
    Args:
        docx_source: Either a file path (str), bytes, or BytesIO object
        
    Returns:
        python-docx Document object
    """
    try:
        if isinstance(docx_source, str):
            return Document(docx_source)
        elif isinstance(docx_source, bytes):
            return Document(BytesIO(docx_source))
        elif isinstance(docx_source, BytesIO):
            docx_source.seek(0)
            return Document(docx_source)
        else:
            raise ValueError(f"Unsupported DOCX source type: {type(docx_source)}")
    
    except Exception as e:
        raise Exception(f"Error parsing DOCX template: {str(e)}")


def get_template_text(doc: Document) -> str:
    """
    Extract all text content from a DOCX document including tables.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        All text content from the document
    """
    text_parts = []
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n".join(text_parts)


def extract_placeholders(doc: Document) -> List[str]:
    """
    Identify placeholder patterns in the document.
    Looks for patterns like {{field}}, [FIELD], <field>, etc.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        List of unique placeholder strings found
    """
    patterns = [
        r'\{\{([^}]+)\}\}',  # {{field}}
        r'\[([A-Z_\s]+)\]',   # [FIELD NAME]
        r'<([^>]+)>',         # <field>
        r'___+',              # Blank lines (underscores)
    ]
    
    all_text = get_template_text(doc)
    placeholders = set()
    
    for pattern in patterns:
        matches = re.findall(pattern, all_text)
        placeholders.update(matches)
    
    return list(placeholders)


def get_document_structure(doc: Document) -> Dict:
    """
    Analyze the document structure to understand its format.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        Dictionary with document structure information
    """
    structure = {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "has_headers": False,
        "table_info": []
    }
    
    # Check for headers
    for section in doc.sections:
        if section.header and section.header.paragraphs:
            for para in section.header.paragraphs:
                if para.text.strip():
                    structure["has_headers"] = True
                    break
    
    # Get table dimensions
    for i, table in enumerate(doc.tables):
        structure["table_info"].append({
            "index": i,
            "rows": len(table.rows),
            "cols": len(table.columns) if table.rows else 0
        })
    
    return structure


if __name__ == "__main__":
    # Test with a sample DOCX
    import sys
    if len(sys.argv) > 1:
        doc = parse_template(sys.argv[1])
        print("Template Text:")
        print(get_template_text(doc))
        print("\nPlaceholders Found:")
        print(extract_placeholders(doc))
